"""
Prometheus (The Illuminator)
============================
The titan who brought fire to mankind - extracts knowledge from documents.

Responsibilities:
- Extract text from PDFs using PyPDF2/pdfplumber
- Extract text from images using pytesseract OCR
- Extract text from DOCX using python-docx
- Extract text from XLSX/CSV using pandas
- Handle multiple file types and return structured content
- Support for multi-page documents
- Image preprocessing for better OCR

Input: List of attached document file paths or file objects
Output: Structured extracted text with metadata (page numbers, sections, etc.)
"""

import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass, asdict
import io
from loguru import logger

# Document processing imports
try:
    import PyPDF2
    import pdfplumber
    from PIL import Image
    import pytesseract
    from docx import Document as DocxDocument
    import pandas as pd
except ImportError as e:
    logger.error(f"Missing document processing library: {e}")


@dataclass
class ExtractedContent:
    """Structured extracted content from a document"""
    filename: str
    file_type: str
    total_pages: int
    content: List[Dict[str, Any]]  # List of {page: int, text: str, metadata: dict}
    raw_text: str
    extraction_method: str
    success: bool
    error: Optional[str] = None
    
    def to_dict(self):
        """Convert to dictionary"""
        return asdict(self)


class PrometheusReader:
    """
    Prometheus (The Illuminator) - Document Text Extraction Agent
    
    Brings the fire of knowledge by extracting text from various document formats.
    """
    
    AGENT_NAME = "Prometheus"
    AGENT_TITLE = "The Illuminator"
    AGENT_DESCRIPTION = "Titan who brought fire to mankind - I extract knowledge from your documents"
    
    SUPPORTED_EXTENSIONS = {
        'pdf': ['.pdf'],
        'image': ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'],
        'document': ['.docx', '.doc'],
        'spreadsheet': ['.xlsx', '.xls', '.csv'],
        'text': ['.txt', '.md', '.json', '.xml', '.html']
    }
    
    def __init__(self, config: Optional[Dict] = None):
        self.name = f"{self.AGENT_NAME} ({self.AGENT_TITLE})"
        self.config = config or {}
        self.ocr_language = self.config.get('ocr_language', 'eng')
        logger.info(f"🔥 {self.name} initialized - Ready to illuminate documents")
    
    async def extract(
        self, 
        file_paths: List[Union[str, Path]], 
        file_contents: Optional[List[bytes]] = None
    ) -> List[ExtractedContent]:
        """
        Extract text from multiple documents.
        
        Args:
            file_paths: List of file paths or filenames
            file_contents: Optional list of file bytes (for uploaded files)
            
        Returns:
            List of ExtractedContent objects
        """
        results = []
        
        for i, file_path in enumerate(file_paths):
            path = Path(file_path)
            file_type = self._get_file_type(path.suffix.lower())
            
            try:
                content_bytes = file_contents[i] if file_contents and i < len(file_contents) else None
                
                if file_type == 'pdf':
                    extracted = await self._extract_pdf(path, content_bytes)
                elif file_type == 'image':
                    extracted = await self._extract_image(path, content_bytes)
                elif file_type == 'document':
                    extracted = await self._extract_docx(path, content_bytes)
                elif file_type == 'spreadsheet':
                    extracted = await self._extract_spreadsheet(path, content_bytes)
                elif file_type == 'text':
                    extracted = await self._extract_text(path, content_bytes)
                else:
                    extracted = ExtractedContent(
                        filename=path.name,
                        file_type='unknown',
                        total_pages=0,
                        content=[],
                        raw_text='',
                        extraction_method='none',
                        success=False,
                        error=f"Unsupported file type: {path.suffix}"
                    )
                
                results.append(extracted)
                logger.info(f"🔥 {self.name}: Extracted {len(extracted.raw_text)} chars from {path.name}")
                
            except Exception as e:
                logger.error(f"🔥 {self.name}: Error extracting {path.name}: {e}")
                results.append(ExtractedContent(
                    filename=path.name,
                    file_type=file_type,
                    total_pages=0,
                    content=[],
                    raw_text='',
                    extraction_method='failed',
                    success=False,
                    error=str(e)
                ))
        
        return results
    
    def _get_file_type(self, extension: str) -> str:
        """Determine file type from extension"""
        for file_type, extensions in self.SUPPORTED_EXTENSIONS.items():
            if extension in extensions:
                return file_type
        return 'unknown'
    
    async def _extract_pdf(self, path: Path, content_bytes: Optional[bytes] = None) -> ExtractedContent:
        """Extract text from PDF using pdfplumber with PyPDF2 fallback"""
        try:
            # Use pdfplumber for better table and layout extraction
            if content_bytes:
                pdf_file = io.BytesIO(content_bytes)
            else:
                pdf_file = str(path)
            
            all_text = []
            content_pages = []
            
            with pdfplumber.open(pdf_file) as pdf:
                total_pages = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    all_text.append(page_text)
                    
                    content_pages.append({
                        "page": page_num,
                        "text": page_text,
                        "metadata": {
                            "width": page.width,
                            "height": page.height
                        }
                    })
            
            raw_text = "\n\n".join(all_text)
            
            return ExtractedContent(
                filename=path.name,
                file_type='pdf',
                total_pages=total_pages,
                content=content_pages,
                raw_text=raw_text,
                extraction_method='pdfplumber',
                success=True
            )
            
        except Exception as e:
            logger.warning(f"pdfplumber failed for {path.name}, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                if content_bytes:
                    pdf_file = io.BytesIO(content_bytes)
                else:
                    pdf_file = open(path, 'rb')
                
                reader = PyPDF2.PdfReader(pdf_file)
                total_pages = len(reader.pages)
                
                all_text = []
                content_pages = []
                
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text() or ""
                    all_text.append(page_text)
                    
                    content_pages.append({
                        "page": page_num,
                        "text": page_text,
                        "metadata": {}
                    })
                
                if not content_bytes:
                    pdf_file.close()
                
                raw_text = "\n\n".join(all_text)
                
                return ExtractedContent(
                    filename=path.name,
                    file_type='pdf',
                    total_pages=total_pages,
                    content=content_pages,
                    raw_text=raw_text,
                    extraction_method='PyPDF2',
                    success=True
                )
                
            except Exception as e2:
                return ExtractedContent(
                    filename=path.name,
                    file_type='pdf',
                    total_pages=0,
                    content=[],
                    raw_text='',
                    extraction_method='failed',
                    success=False,
                    error=f"Both pdfplumber and PyPDF2 failed: {e2}"
                )
    
    async def _extract_image(self, path: Path, content_bytes: Optional[bytes] = None) -> ExtractedContent:
        """Extract text from image using OCR"""
        try:
            if content_bytes:
                image = Image.open(io.BytesIO(content_bytes))
            else:
                image = Image.open(path)
            
            # Perform OCR
            text = pytesseract.image_to_string(image, lang=self.ocr_language)
            
            return ExtractedContent(
                filename=path.name,
                file_type='image',
                total_pages=1,
                content=[{
                    "page": 1,
                    "text": text,
                    "metadata": {
                        "width": image.width,
                        "height": image.height,
                        "format": image.format
                    }
                }],
                raw_text=text,
                extraction_method='pytesseract',
                success=True
            )
            
        except Exception as e:
            return ExtractedContent(
                filename=path.name,
                file_type='image',
                total_pages=0,
                content=[],
                raw_text='',
                extraction_method='failed',
                success=False,
                error=f"OCR extraction failed: {e}"
            )
    
    async def _extract_docx(self, path: Path, content_bytes: Optional[bytes] = None) -> ExtractedContent:
        """Extract text from DOCX"""
        try:
            if content_bytes:
                doc = DocxDocument(io.BytesIO(content_bytes))
            else:
                doc = DocxDocument(path)
            
            # Extract all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            raw_text = "\n\n".join(paragraphs)
            
            return ExtractedContent(
                filename=path.name,
                file_type='docx',
                total_pages=1,  # DOCX doesn't have strict pages
                content=[{
                    "page": 1,
                    "text": raw_text,
                    "metadata": {
                        "paragraph_count": len(paragraphs)
                    }
                }],
                raw_text=raw_text,
                extraction_method='python-docx',
                success=True
            )
            
        except Exception as e:
            return ExtractedContent(
                filename=path.name,
                file_type='docx',
                total_pages=0,
                content=[],
                raw_text='',
                extraction_method='failed',
                success=False,
                error=f"DOCX extraction failed: {e}"
            )
    
    async def _extract_spreadsheet(self, path: Path, content_bytes: Optional[bytes] = None) -> ExtractedContent:
        """Extract text from spreadsheet"""
        try:
            if content_bytes:
                if path.suffix.lower() == '.csv':
                    df = pd.read_csv(io.BytesIO(content_bytes))
                else:
                    df = pd.read_excel(io.BytesIO(content_bytes))
            else:
                if path.suffix.lower() == '.csv':
                    df = pd.read_csv(path)
                else:
                    df = pd.read_excel(path)
            
            # Convert to string representation
            text = df.to_string(index=False)
            
            return ExtractedContent(
                filename=path.name,
                file_type='spreadsheet',
                total_pages=1,
                content=[{
                    "page": 1,
                    "text": text,
                    "metadata": {
                        "rows": len(df),
                        "columns": len(df.columns),
                        "column_names": list(df.columns)
                    }
                }],
                raw_text=text,
                extraction_method='pandas',
                success=True
            )
            
        except Exception as e:
            return ExtractedContent(
                filename=path.name,
                file_type='spreadsheet',
                total_pages=0,
                content=[],
                raw_text='',
                extraction_method='failed',
                success=False,
                error=f"Spreadsheet extraction failed: {e}"
            )
    
    async def _extract_text(self, path: Path, content_bytes: Optional[bytes] = None) -> ExtractedContent:
        """Extract text from plain text files"""
        try:
            if content_bytes:
                text = content_bytes.decode('utf-8', errors='ignore')
            else:
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            
            return ExtractedContent(
                filename=path.name,
                file_type='text',
                total_pages=1,
                content=[{
                    "page": 1,
                    "text": text,
                    "metadata": {
                        "char_count": len(text),
                        "line_count": text.count('\n') + 1
                    }
                }],
                raw_text=text,
                extraction_method='direct',
                success=True
            )
            
        except Exception as e:
            return ExtractedContent(
                filename=path.name,
                file_type='text',
                total_pages=0,
                content=[],
                raw_text='',
                extraction_method='failed',
                success=False,
                error=f"Text extraction failed: {e}"
            )
    
    def get_agent_info(self) -> Dict[str, str]:
        """Return agent information for UI display"""
        return {
            "name": self.AGENT_NAME,
            "title": self.AGENT_TITLE,
            "full_name": self.name,
            "description": self.AGENT_DESCRIPTION,
            "icon": "🔥",
            "role": "document_reader"
        }
